import docx
from tempgen.parsers.parser import AbstractParser

class Parser(AbstractParser):
    def paragraph_replace_text(self, paragraph, str, replace_str):
        '''
        https://github.com/python-openxml/python-docx/issues/30#issuecomment-881106471
        '''
        count = 0
        search_pos = 0
        while paragraph.text.find(str, search_pos) != -1:
            match = { 'start': paragraph.text.find(str, search_pos), 'end': paragraph.text.find(str, search_pos) + len(str) }
            search_pos = match['end']
            padding = (len(replace_str) - (match['end'] -match['start']) ) *count
            runs = iter(paragraph.runs)
            start, end = match['start'] + padding , match['end'] + padding
            for run in runs:
                run_len = len(run.text)
                if start < run_len:
                    break
                start, end = start - run_len, end - run_len
            run_text = run.text
            run_len = len(run_text)
            run.text = '%s%s%s' % (run_text[:start], replace_str, run_text[end:])
            end -= run_len
            for run in runs:
                if end <= 0:
                    break
                run_text = run.text
                run_len = len(run_text)
                run.text = run_text[end:]
                end -= run_len
            count += 1
        return paragraph

    def replace_in_paragraph(self, p, d):
        for replaced, replacement in d.items():
            self.paragraph_replace_text(p, replaced, replacement)

    def collect_paragraphs(self, doc):
        paragraphs = []
        for p in doc.paragraphs:
                paragraphs.append(p)
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for p in cell.paragraphs:
                        paragraphs.append(p)
        return paragraphs

    def parse(self, path, container, parse_entry, find_matches):
        doc = docx.Document(path)
        paragraphs = self.collect_paragraphs(doc)
        for p in paragraphs:
            matches = find_matches(p.text)
            for match in matches:
                payload = parse_entry(match, path)
                container[payload['id']] = payload

    def replace(self, source_path, target_path, compute_match, replacements, update_external = False):
        doc = docx.Document(source_path)
        paragraphs = self.collect_paragraphs(doc)
        to_replace = {}
        for p in paragraphs:
            compute_match(p.text, to_replace, replacements, source_path, update_external)
            self.replace_in_paragraph(p, to_replace)
        doc.save(target_path)